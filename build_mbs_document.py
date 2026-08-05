from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path


OUT = Path(r"C:\Users\marce\Documents\Planilhas Panda\MBS Ofertas - Documento Completo do Projeto.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F7F9FC"
GOLD = "9A6A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111111"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[i] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    if keep_lines:
        keep = OxmlElement("w:keepLines")
        p_pr.append(keep)


def set_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(r_pr)
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_numbering_definition(doc, fmt, text, left=540, hanging=270):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(tabs)
    p_pr.append(ind)
    p_pr.append(spacing)
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

code_style = styles.add_style("MBS Code", WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = "Consolas"
code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor.from_string(NAVY)
code_style.paragraph_format.space_before = Pt(3)
code_style.paragraph_format.space_after = Pt(5)
code_style.paragraph_format.line_spacing = 1.08

bullet_num = add_numbering_definition(doc, "bullet", "•")
decimal_num = add_numbering_definition(doc, "decimal", "%1.")

# Running header/footer
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_run = header_p.add_run("MBS OFERTAS  |  DOCUMENTO MESTRE DO PROJETO")
set_font(header_run, size=8.5, color=MUTED, bold=True)
footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer_p.add_run("MBS Ofertas  •  Página ")
set_font(footer_run, size=9, color=MUTED)
add_page_number(footer_p)

# Cover - editorial_cover pattern
for _ in range(4):
    doc.add_paragraph()
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = kicker.add_run("DOCUMENTAÇÃO DO PRODUTO")
set_font(r, size=10.5, color=GOLD, bold=True)
kicker.paragraph_format.space_after = Pt(18)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("MBS Ofertas")
set_font(r, size=31, color=NAVY, bold=True)
title.paragraph_format.space_after = Pt(8)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Conversor de anúncios da Amazon para divulgação")
set_font(r, size=15, color=DARK_BLUE)
subtitle.paragraph_format.space_after = Pt(4)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Especificação funcional, regras, configuração do GPT e histórico disponível")
set_font(r, size=11.5, color=MUTED, italic=True)
sub2.paragraph_format.space_after = Pt(76)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Versão 1.0  |  31 de julho de 2026")
set_font(r, size=10, color=MUTED, bold=True)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run("Baseado no conteúdo desta tarefa e na análise do site publicado.")
set_font(r, size=9.5, color=MUTED)
doc.add_page_break()


def add_heading(text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_repeat_keep(p, keep_next=True)
    return p


def add_body(text, bold_lead=None, italic=False):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, italic=italic)
    return p


def add_bullet(text):
    p = doc.add_paragraph()
    apply_numbering(p, bullet_num)
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(text):
    p = doc.add_paragraph()
    apply_numbering(p, decimal_num)
    r = p.add_run(text)
    set_font(r)
    return p


def add_callout(label, text, fill=LIGHT_BLUE, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r = p.add_run(label + " ")
    set_font(r, color=color, bold=True)
    r = p.add_run(text)
    set_font(r, color=BLACK)


add_heading("1. Resumo executivo", 1)
add_body("O MBS Ofertas é um conversor de anúncios promocionais da Amazon Brasil. Ele recebe textos normalmente encaminhados por WhatsApp, identifica produtos e links, substitui os endereços por links vinculados ao identificador de afiliado MBS e devolve mensagens prontas para compartilhamento.")
add_callout("Decisão central:", "o GPT pessoal deve usar o modo “Trocar links + reformular” como padrão e o identificador de afiliado mbs19770a-20 em todos os links Amazon convertidos.")
add_body("O projeto reduz trabalho manual, padroniza a divulgação e diminui erros na troca de links. Nenhuma senha ou conexão com a conta Amazon é necessária; o processamento utiliza somente os textos e links fornecidos pelo usuário.")

add_heading("2. Escopo e origem deste documento", 1)
add_body("Este documento consolida integralmente as informações disponíveis nesta tarefa: a solicitação para detalhar o projeto, a inspeção do site publicado e a configuração proposta para um GPT pessoal. Ele foi criado para servir como memória do produto e como base para manutenção, evolução ou reconstrução do MBS Ofertas.")
add_callout("Limite do histórico:", "o assistente não possui acesso automático a outros bate-papos que não estejam presentes nesta tarefa. Conversas antigas poderão ser incorporadas em uma versão futura se forem copiadas ou exportadas pelo usuário.", fill="FFF7E0", color=GOLD)

add_heading("3. Estado observado do projeto", 1)
add_body("Endereço analisado:")
p = doc.add_paragraph()
add_hyperlink(p, "Abrir o site publicado do MBS Ofertas", "https://conversor-ofertas-amazon.jufap.chatgpt.site/")
add_body("Identificador de afiliado exibido: mbs19770a-20.")
add_body("A interface publicada apresenta três etapas: colar os anúncios, escolher o tipo de conversão e copiar os anúncios prontos. Ela informa a quantidade de caracteres, a quantidade de links encontrados, o total de ofertas processadas e o status individual de cada link.")

add_heading("3.1 Elementos da interface", 2)
for item in [
    "Área de texto para colar um ou vários anúncios recebidos no WhatsApp.",
    "Botão “Usar exemplo” para preencher uma oferta demonstrativa.",
    "Modo “Trocar links + reformular”, marcado como recomendado e selecionado por padrão.",
    "Modo “Somente trocar links”, que preserva o texto original.",
    "Botão “Converter anúncios”.",
    "Área de resultados com status “Link convertido” ou “Verifique o link”.",
    "Botões para copiar cada anúncio ou copiar todos.",
    "Aviso de que preços e condições podem mudar e devem ser conferidos antes do compartilhamento.",
]:
    add_bullet(item)

add_heading("3.2 Observação de qualidade", 2)
add_body("No exemplo publicado, uma linha com cinco hífens foi interpretada como separador e produziu duas ofertas. A primeira ficou sem link e recebeu o status “Verifique o link”. Para o GPT pessoal, a regra foi ajustada: o separador principal deve ser uma linha isolada contendo exatamente três hífens, e divisões ambíguas devem preservar o anúncio como uma única oferta.")

add_heading("4. Jornada do usuário", 1)
for step in [
    "O usuário cola uma ou várias mensagens promocionais.",
    "O sistema identifica os limites de cada oferta.",
    "O usuário escolhe entre reformular o anúncio ou somente trocar os links.",
    "O sistema identifica todos os links Amazon e resolve os links encurtados quando possível.",
    "O sistema extrai o ASIN e gera o link afiliado limpo.",
    "O texto é preservado ou reformulado conforme o modo escolhido.",
    "Cada anúncio recebe um status e o aviso obrigatório de variação de preço e condições.",
    "O resultado é entregue pronto para copiar e compartilhar.",
]:
    add_number(step)

add_heading("5. Requisitos funcionais", 1)
for item in [
    "Aceitar texto simples, emojis e formatação de WhatsApp.",
    "Processar uma ou várias ofertas na mesma entrada.",
    "Detectar amazon.com.br, www.amazon.com.br, amzn.to, a.co e formatos equivalentes.",
    "Converter todos os links Amazon encontrados em cada oferta.",
    "Preservar links que não sejam da Amazon.",
    "Remover identificadores de afiliado anteriores e aplicar mbs19770a-20.",
    "Nunca inventar ASIN, preço, cupom, estoque ou condição comercial.",
    "Sinalizar links que não puderem ser resolvidos.",
    "Entregar os anúncios em blocos independentes e prontos para copiar.",
]:
    add_bullet(item)

add_heading("6. Modos de conversão", 1)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
set_table_geometry(table, [2100, 4260, 3000])
headers = ["Modo", "Comportamento", "Quando usar"]
for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=10, color=NAVY, bold=True)
set_repeat_table_header(table.rows[0])
rows = [
    ("Trocar links + reformular", "Converte os links e reorganiza a mensagem, mantendo todos os fatos comerciais.", "Uso padrão para criar uma apresentação própria."),
    ("Somente trocar links", "Mantém o texto original e altera apenas os links Amazon.", "Quando a mensagem já está aprovada ou deve permanecer idêntica."),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        p = cells[i].paragraphs[0]
        r = p.add_run(text)
        set_font(r, size=9.5, bold=(i == 0))
set_table_geometry(table, [2100, 4260, 3000])

add_heading("7. Regras de conversão de links", 1)
add_heading("7.1 Links completos", 2)
for step in [
    "Identificar o ASIN nos formatos /dp/ASIN, /gp/product/ASIN ou equivalente.",
    "Descartar parâmetros de rastreamento e identificadores de afiliado anteriores.",
    "Gerar o endereço canônico com o identificador MBS.",
]:
    add_number(step)
p = doc.add_paragraph(style="MBS Code")
p.add_run("https://www.amazon.com.br/dp/ASIN?tag=mbs19770a-20")

add_heading("7.2 Links encurtados", 2)
for step in [
    "Abrir o endereço amzn.to ou a.co usando navegação na web.",
    "Observar o endereço final do produto.",
    "Extrair o ASIN sem adivinhar.",
    "Gerar o link afiliado canônico.",
]:
    add_number(step)
add_body("Se o endereço não puder ser aberto ou o ASIN não for identificado, o link original deve ser mantido com o aviso: “Link não convertido — verifique manualmente”.")

add_heading("7.3 Casos especiais", 2)
for item in [
    "Vários links Amazon em uma oferta: converter todos.",
    "Link que já contém outro tag: substituir pelo identificador MBS.",
    "Link não Amazon: preservar sem alteração.",
    "Link quebrado ou inacessível: não fabricar destino; sinalizar a necessidade de conferência.",
    "Produto com variação: não afirmar que cor, tamanho ou versão do anúncio corresponde à página sem verificação.",
]:
    add_bullet(item)

add_heading("8. Regras de reformulação", 1)
add_body("A reformulação deve melhorar a leitura e dar identidade própria ao anúncio sem alterar seus fatos. O texto pode ser reorganizado, mas não pode ganhar informações que não estavam na entrada ou na página verificada.")
add_heading("8.1 Informações obrigatoriamente preservadas", 2)
for item in [
    "Nome e versão exata do produto.",
    "Preço anunciado e valor normal, quando informados.",
    "Cupom e instruções para aplicá-lo.",
    "Quantidade mínima exigida.",
    "Frete, parcelamento e formas de pagamento.",
    "Tamanho, cor, volume, modelo e demais variações.",
    "Observações essenciais sobre a condição da promoção.",
]:
    add_bullet(item)

add_heading("8.2 Alterações permitidas", 2)
for item in [
    "Reorganizar parágrafos, emojis e chamadas promocionais.",
    "Corrigir pequenos erros de ortografia e espaçamento.",
    "Reduzir repetições exageradas.",
    "Variar introduções para que os anúncios não pareçam idênticos.",
    "Colocar o link convertido em uma linha própria próximo ao final.",
]:
    add_bullet(item)

add_heading("8.3 Conteúdo proibido", 2)
for item in [
    "Inventar preço, desconto, cupom, frete grátis ou estoque.",
    "Inventar avaliações, benefícios ou especificações.",
    "Criar urgência que não esteja informada.",
    "Afirmar que a promoção ainda está ativa sem verificação atual.",
    "Alterar silenciosamente uma informação ambígua.",
]:
    add_bullet(item)

add_heading("9. Formato de saída", 1)
add_body("A resposta padrão começa com um resumo operacional e depois apresenta cada anúncio em bloco separado.")
p = doc.add_paragraph(style="MBS Code")
p.add_run("[OK] X ofertas processadas\n[LINK] X links convertidos\n[ATENÇÃO] X links para verificar\n\nOFERTA 01 — Link convertido\n\n[anúncio completo pronto para copiar]\n\n---\n\nOFERTA 02 — Verifique o link\n\n[anúncio completo pronto para copiar]")
add_body("Aviso obrigatório ao final de cada anúncio:")
p = doc.add_paragraph(style="MBS Code")
p.add_run("⚠️ Confira o preço, o frete e as condições antes de finalizar, pois podem mudar.")

add_heading("10. Configuração recomendada do GPT pessoal", 1)
add_body("Nome: MBS Ofertas", bold_lead="Nome:")
add_body("Descrição: Converte anúncios da Amazon em mensagens prontas para WhatsApp, substituindo os links pelo identificador de afiliado MBS e, opcionalmente, reformulando a apresentação da oferta.", bold_lead="Descrição:")
add_body("Capacidade recomendada: pesquisa ou navegação na web ativada. Essa capacidade é necessária para resolver links encurtados e verificar páginas quando solicitado.", bold_lead="Capacidade recomendada:")

add_heading("10.1 Iniciadores de conversa", 2)
for item in [
    "Converta estas ofertas e reformule os textos.",
    "Troque apenas os links Amazon destes anúncios.",
    "Organize estas promoções para divulgar no WhatsApp.",
    "Verifique quais links não puderam ser convertidos.",
]:
    add_bullet(item)

add_heading("11. Instruções completas para colar no GPT", 1)
add_callout("Uso:", "copie todo o conteúdo desta seção e cole no campo Instruções do seu GPT pessoal.")

instruction_sections = [
    ("IDENTIDADE", [
        "Você é o MBS Ofertas, um assistente especializado em converter anúncios de produtos da Amazon Brasil em mensagens prontas para divulgação no WhatsApp.",
        "Use sempre o identificador de afiliado mbs19770a-20.",
    ]),
    ("OBJETIVO", [
        "Receber um ou vários anúncios, identificar os links da Amazon, convertê-los para links de afiliado e devolver cada oferta pronta para copiar e compartilhar.",
    ]),
    ("MODO PADRÃO", [
        "Quando o usuário apenas colar os anúncios sem dar outra instrução, use automaticamente o modo Trocar links + reformular.",
        "No modo Trocar links + reformular, substitua os links Amazon e reorganize o texto para criar uma apresentação própria, clara e atraente.",
        "No modo Somente trocar links, mantenha rigorosamente o texto original, modifique somente os links da Amazon e acrescente o aviso obrigatório ao final.",
    ]),
    ("SEPARAÇÃO DAS OFERTAS", [
        "Considere uma linha contendo somente --- como separador entre anúncios.",
        "Não divida anúncios por travessões usados em preços, listas ou elementos decorativos. Linhas como ----- só devem ser consideradas separadores quando estiver claro que foram usadas deliberadamente entre dois anúncios completos.",
        "Se o conteúdo puder representar uma única oferta, preserve-o como uma única oferta.",
    ]),
    ("CONVERSÃO DOS LINKS", [
        "Identifique links amazon.com.br, www.amazon.com.br, amzn.to, a.co, links com parâmetros de afiliado e links nos formatos /dp/ASIN ou /gp/product/ASIN.",
        "Para links completos, identifique o ASIN, remova identificadores de afiliado anteriores e gere https://www.amazon.com.br/dp/ASIN?tag=mbs19770a-20.",
        "Para links encurtados, abra o link usando navegação na web, localize o endereço final, extraia o ASIN e gere o link limpo com o identificador MBS.",
        "Nunca invente um ASIN. Se não conseguir abrir ou identificar o produto, mantenha o link original e sinalize: Link não convertido — verifique manualmente.",
        "Quando houver vários links Amazon na mesma oferta, converta todos. Preserve links que não sejam da Amazon.",
    ]),
    ("REFORMULAÇÃO", [
        "Preserve o nome e a versão exata do produto, preço, desconto, cupom, quantidade mínima, pagamento, parcelas, frete, tamanho, cor, volume, modelo e outras condições.",
        "Preserve a formatação de WhatsApp, incluindo *negrito* e _itálico_.",
        "Pode reorganizar parágrafos, emojis e chamadas, corrigir pequenos erros e reduzir repetições exageradas.",
        "Use linguagem promocional, direta e natural. Varie as introduções e coloque o link convertido em linha própria próximo ao final.",
        "Não invente preços, descontos, cupons, frete grátis, estoque, avaliações, benefícios, urgência ou especificações.",
        "Não afirme que uma oferta continua ativa sem verificá-la naquele momento.",
    ]),
    ("AVISO OBRIGATÓRIO", [
        "Acrescente ao final de cada anúncio: ⚠️ Confira o preço, o frete e as condições antes de finalizar, pois podem mudar.",
    ]),
    ("FORMATO DA RESPOSTA", [
        "Comece com um resumo: X ofertas processadas, X links convertidos e X links para verificar.",
        "Depois apresente cada oferta separadamente, com numeração e status Link convertido ou Verifique o link.",
        "Não coloque comentários ou explicações dentro do anúncio.",
        "Quando todos os links forem convertidos corretamente, não faça perguntas adicionais; entregue o resultado.",
        "Quando uma informação estiver ambígua, preserve o texto original em vez de adivinhar.",
    ]),
    ("PEDIDOS COMPLEMENTARES", [
        "Se o usuário pedir somente os anúncios, omita o resumo e entregue apenas os textos finais.",
        "Se pedir uma versão curta, reduza repetições, mas preserve as condições essenciais.",
        "Se pedir uma versão para WhatsApp, mantenha links visíveis e formatação compatível.",
        "Se pedir validação de preço, abra as páginas e informe separadamente o que foi possível confirmar, sem alterar silenciosamente os valores enviados.",
    ]),
    ("PRIVACIDADE", [
        "Nunca solicite senha, acesso à conta Amazon ou dados pessoais. A conversão funciona somente com textos e links enviados pelo usuário.",
    ]),
]

for label, paragraphs in instruction_sections:
    p = doc.add_paragraph(style="MBS Code")
    r = p.add_run(label)
    r.bold = True
    for text in paragraphs:
        p = doc.add_paragraph(style="MBS Code")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run("• " + text)

add_heading("12. Casos de teste", 1)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
set_table_geometry(table, [2400, 3580, 3380])
for i, text in enumerate(["Cenário", "Entrada", "Resultado esperado"]):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, size=10, color=NAVY, bold=True)
set_repeat_table_header(table.rows[0])
tests = [
    ("Link completo", "URL amazon.com.br/dp/ASIN com outro tag", "Link canônico com tag mbs19770a-20."),
    ("Link encurtado", "URL amzn.to válida", "Redirecionamento resolvido, ASIN extraído e link afiliado gerado."),
    ("Link inacessível", "URL encurtada quebrada", "Link original preservado e aviso para conferir."),
    ("Vários produtos", "Três anúncios separados por linha ---", "Três ofertas independentes, todas numeradas."),
    ("Separador decorativo", "Cinco hífens dentro de uma única mensagem", "Mensagem preservada quando não houver dois anúncios completos."),
    ("Somente links", "Pedido explícito para não reescrever", "Texto idêntico, exceto links e aviso obrigatório."),
    ("Reformulação", "Anúncio repetitivo com preço e cupom", "Texto mais claro sem perda ou criação de fatos."),
]
for row in tests:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        p = cells[i].paragraphs[0]
        r = p.add_run(text)
        set_font(r, size=9.2, bold=(i == 0))
set_table_geometry(table, [2400, 3580, 3380])

add_heading("13. Critérios de aceitação", 1)
for item in [
    "Todos os links Amazon válidos recebem o identificador mbs19770a-20.",
    "Nenhum ASIN é criado por inferência ou adivinhação.",
    "Links não Amazon permanecem intactos.",
    "Informações comerciais são preservadas integralmente.",
    "O modo Somente trocar links não reescreve a mensagem.",
    "Cada oferta recebe status claro e aviso obrigatório.",
    "O resultado pode ser copiado diretamente para o WhatsApp.",
    "A separação em lote não fragmenta uma única oferta por engano.",
]:
    add_bullet(item)

add_heading("14. Limitações e dependências", 1)
for item in [
    "Links encurtados dependem de acesso à web e da disponibilidade do redirecionamento.",
    "Preços, frete, cupons e estoque podem mudar depois do processamento.",
    "Uma instrução textual não substitui validação automática pela API oficial da Amazon.",
    "O GPT deve informar incertezas em vez de preencher informações ausentes.",
    "O histórico deste documento está limitado às conversas disponíveis nesta tarefa.",
]:
    add_bullet(item)

add_heading("15. Melhorias recomendadas", 1)
for item in [
    "Permitir que o identificador de afiliado seja configurado sem alterar as instruções principais.",
    "Validar automaticamente se cada ASIN corresponde ao produto anunciado.",
    "Criar estilos de reformulação: direto, urgente, premium, curto e sem emojis.",
    "Oferecer pré-visualização específica para WhatsApp.",
    "Registrar anúncios processados para evitar repetição de texto.",
    "Adicionar opção de exportação para planilha ou calendário editorial.",
    "Criar uma biblioteca de testes com links completos, encurtados, quebrados e múltiplos produtos.",
]:
    add_bullet(item)

add_heading("16. Registro do bate-papo disponível", 1)
add_body("Solicitação inicial do usuário:", bold_lead="Solicitação inicial do usuário:")
add_callout("Pedido:", "“detalha esse projeto MBS Ofertas, quero colocar no meu GPT pessoal”", fill=PALE)
add_body("Resultado produzido: análise do site, identificação do fluxo e do ID de afiliado, definição dos dois modos de conversão, regras de tratamento de links e um conjunto completo de instruções para o GPT pessoal.")
add_body("Solicitação complementar do usuário:", bold_lead="Solicitação complementar do usuário:")
add_callout("Pedido:", "“adicione também um documento com tudo que foi escrito para criar esse projeto nos nossos bate papos”", fill=PALE)
add_body("Resultado produzido: este documento mestre, que consolida o conteúdo disponível e cria uma base durável para continuidade do projeto.")

add_heading("17. Próximo passo para completar o histórico", 1)
add_body("Para incorporar bate-papos anteriores, exporte ou copie as conversas antigas e envie os arquivos ou textos. A próxima versão poderá acrescentar uma linha do tempo das decisões, prompts originais, alterações solicitadas, versões do produto e justificativas de cada escolha.")

# Metadata
doc.core_properties.title = "MBS Ofertas - Documento Completo do Projeto"
doc.core_properties.subject = "Especificação funcional e configuração do GPT pessoal"
doc.core_properties.author = "MBS Ofertas"
doc.core_properties.keywords = "MBS Ofertas, Amazon, afiliado, GPT, WhatsApp"

doc.save(OUT)
print(str(OUT))
