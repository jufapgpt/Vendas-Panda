from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
from zipfile import ZipFile

path = Path(r"C:\Users\marce\Documents\Planilhas Panda\MBS Ofertas - Documento Completo do Projeto.docx")
doc = Document(path)

required = [
    "Resumo executivo",
    "Estado observado do projeto",
    "Requisitos funcionais",
    "Regras de conversão de links",
    "Configuração recomendada do GPT pessoal",
    "Instruções completas para colar no GPT",
    "Casos de teste",
    "Critérios de aceitação",
    "Registro do bate-papo disponível",
]
all_text = "\n".join(p.text for p in doc.paragraphs)
missing = [item for item in required if item not in all_text]

table_reports = []
for idx, table in enumerate(doc.tables, start=1):
    grid = table._tbl.tblGrid
    widths = [int(col.get(qn("w:w"))) for col in grid.findall(qn("w:gridCol"))]
    row_ok = True
    for row in table.rows:
        cell_widths = []
        for cell in row.cells:
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            cell_widths.append(int(tc_w.get(qn("w:w"))))
        if cell_widths != widths:
            row_ok = False
    table_reports.append({"table": idx, "columns": widths, "sum": sum(widths), "consistent": row_ok})

with ZipFile(path) as archive:
    bad = archive.testzip()
    members = set(archive.namelist())
    essentials = {"word/document.xml", "word/styles.xml", "word/numbering.xml", "[Content_Types].xml"}
    missing_parts = sorted(essentials - members)

print({
    "file": str(path),
    "size_bytes": path.stat().st_size,
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "sections": len(doc.sections),
    "missing_required_sections": missing,
    "zip_bad_member": bad,
    "missing_ooxml_parts": missing_parts,
    "table_geometry": table_reports,
})

if missing or bad or missing_parts or not all(r["consistent"] for r in table_reports):
    raise SystemExit(1)
